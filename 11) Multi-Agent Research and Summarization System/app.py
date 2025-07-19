import os
import streamlit as st
import pdfplumber
from docx import Document as DocxDocument

from langchain_community.vectorstores import FAISS
from langchain.docstore.document import Document
from langchain.chains import RetrievalQA
from langchain.prompts import PromptTemplate
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langgraph.graph import StateGraph
from langchain_core.runnables import RunnableLambda

from langchain_google_genai import ChatGoogleGenerativeAI, GoogleGenerativeAIEmbeddings
from langchain_community.tools import DuckDuckGoSearchRun

# ---------------------- CONFIGURATION ----------------------
GOOGLE_API_KEY = "AIzaSyC7K1mPTvB9WDVC06u31HlkvBzH0hMOdbA"

if not GOOGLE_API_KEY:
    raise ValueError("GOOGLE_API_KEY not found in environment variables")

llm = ChatGoogleGenerativeAI(model="gemini-1.5-flash", temperature=0.3, google_api_key=GOOGLE_API_KEY)
embeddings = GoogleGenerativeAIEmbeddings(model="models/embedding-001", google_api_key=GOOGLE_API_KEY)
text_splitter = RecursiveCharacterTextSplitter(chunk_size=500, chunk_overlap=50)
search = DuckDuckGoSearchRun()

# ---------------------- FILE PARSER ----------------------
def extract_text_from_local_path(path):
    if path.endswith(".pdf"):
        with pdfplumber.open(path) as pdf:
            return "\n".join([page.extract_text() for page in pdf.pages if page.extract_text()])
    elif path.endswith(".txt"):
        with open(path, "r", encoding="utf-8") as f:
            return f.read()
    elif path.endswith(".docx"):
        doc = DocxDocument(path)
        return "\n".join([p.text for p in doc.paragraphs])
    return ""

# ---------------------- AGENTS ----------------------
def router_agent(state):
    query = state.get("query", "")
    route_prompt = PromptTemplate.from_template(
        "Classify the query into one of [web, rag, llm]:\n\nQuery: {query}\n\nAnswer:"
    )
    route_result = (route_prompt | llm).invoke({"query": query}).content.lower()
    route = "llm"
    if "web" in route_result:
        route = "web"
    elif "rag" in route_result:
        route = "rag"
    return {**state, "route": route}

def web_agent(state):
    query = state["query"]
    try:
        result = search.run(query)
        return {**state, "content": result}
    except Exception as e:
        return {**state, "content": f"Web search failed: {str(e)}"}

def rag_agent(state):
    query = state["query"]
    retriever = state["retriever"]
    qa_chain = RetrievalQA.from_chain_type(llm=llm, retriever=retriever)
    answer = qa_chain.run(query)
    return {**state, "content": answer}

def llm_agent(state):
    query = state["query"]
    response = llm.invoke(query)
    return {**state, "content": response.content}

def summarizer_agent(state):
    content = state["content"]
    prompt = PromptTemplate.from_template("Summarize clearly and concisely:\n\n{content}")
    summary = (prompt | llm).invoke({"content": content}).content
    return {**state, "final": summary}

# ---------------------- LANGGRAPH ----------------------
def run_langgraph(user_query, retriever):
    workflow = StateGraph(dict)
    workflow.set_entry_point("router")

    workflow.add_node("router", RunnableLambda(router_agent))
    workflow.add_node("web", RunnableLambda(web_agent))
    workflow.add_node("rag", RunnableLambda(rag_agent))
    workflow.add_node("llm", RunnableLambda(llm_agent))
    workflow.add_node("summarizer", RunnableLambda(summarizer_agent))

    def router_logic(state): return state["route"]
    workflow.add_conditional_edges("router", router_logic, {
        "web": "web",
        "rag": "rag",
        "llm": "llm"
    })

    for node in ["web", "rag", "llm"]:
        workflow.add_edge(node, "summarizer")

    workflow.set_finish_point("summarizer")
    app = workflow.compile()
    return app.invoke({"query": user_query, "retriever": retriever})["final"]

# ---------------------- STREAMLIT APP ----------------------
# Enhanced UI
st.set_page_config(page_title="🔍 Fully Agentic Research Assistant", layout="wide", page_icon="🧠")

# Sidebar for branding and info
with st.sidebar:
    st.image("https://img.icons8.com/color/96/000000/artificial-intelligence.png", width=80)
    st.markdown("""
    # 🧠 Multi-Agent RAG System
    **LangGraph + Web + RAG + LLM**
    
    <hr style='border:1px solid #eee'>
    
    - [GitHub](https://github.com/VikramS75/Agentic_AI_Workshop/tree/main/11%29%20Multi-Agent%20Research%20and%20Summarization%20System)
    - [LangGraph Docs](https://langchain-ai.github.io/langgraph/)
    
    <hr style='border:1px solid #eee'>
    
    <small>Powered by Gemini, DuckDuckGo, FAISS, LangChain</small>
    """, unsafe_allow_html=True)

st.markdown("""
<style>
.big-title { font-size: 2.5rem; font-weight: bold; color: #4F8BF9; }
.subtext { color: #666; font-size: 1.1rem; }
.question-box { background: #f7fafd; border-radius: 10px; padding: 1.5em; margin-bottom: 1em; border: 1px solid #e3eafc; }
.answer-box { background: #eaf6ef; border-radius: 10px; padding: 1.5em; border: 1px solid #b7e4c7; }
</style>
""", unsafe_allow_html=True)

st.markdown('<div class="big-title">🧠 Multi-Agent Research & Summarization System</div>', unsafe_allow_html=True)
st.markdown('<div class="subtext">Ask anything! The system will smartly route your query to the best agent: <b>Web Search</b>, <b>RAG</b>, or <b>LLM</b>. Results are always summarized for you.</div>', unsafe_allow_html=True)
st.markdown("---")

retriever = None
documents_loaded = False

# Load local documents
if os.path.exists("my_docs"):
    with st.spinner("📂 Loading documents from 'my_docs' folder..."):
        all_content = []
        for filename in os.listdir("my_docs"):
            filepath = os.path.join("my_docs", filename)
            if filename.lower().endswith(('.pdf', '.txt', '.docx')):
                content = extract_text_from_local_path(filepath)
                if content:
                    all_content.append(content)

        if all_content:
            chunks = text_splitter.create_documents(all_content)
            vectorstore = FAISS.from_documents(chunks, embeddings)
            retriever = vectorstore.as_retriever()
            documents_loaded = True
            st.success(f"✅ Loaded {len(all_content)} documents.")
        else:
            st.warning("⚠️ No readable files found.")

if not documents_loaded:
    st.info("📄 Using fallback knowledge base.")
    docs = [
        Document(page_content="LangGraph is a Python framework for agent workflows."),
        Document(page_content="Gemini 1.5 Flash is fast and great for summarization."),
    ]
    vectorstore = FAISS.from_documents(docs, embeddings)
    retriever = vectorstore.as_retriever()

# User Input
center_col1, center_col2, center_col3 = st.columns([2, 3, 2])
with center_col2:
    # st.markdown('<div class="question-box">', unsafe_allow_html=True)
    query = st.text_input("💬 Ask your question", placeholder="e.g. What is LangGraph?", key="user_query")
    submit = st.button("🚀 Submit", use_container_width=True)
    st.markdown('</div>', unsafe_allow_html=True)

if submit:
    if not query.strip():
        st.warning("⚠️ Please enter a question.")
    else:
        with st.spinner("🤖 Thinking..."):
            try:
                answer = run_langgraph(query, retriever)
                st.success("✅ Done!")
                with center_col2:
                    st.markdown('<div class="answer-box">', unsafe_allow_html=True)
                    st.subheader("📘 Answer:")
                    st.write(answer)
                    st.markdown('</div>', unsafe_allow_html=True)
            except Exception as e:
                st.error(f"❌ Error: {str(e)}")